import PyPDF2
import re
import os
import translator
import docx
import summarizer


class pdfParser(object):
    """docstring for pdfParser"""

    def __init__(self, filehandle=None):
        super(pdfParser, self).__init__()
        if filehandle:
            self.open(filehandle)
        if not os.path.isdir('tmp'):
            os.mkdir('tmp')
        self.tempPath = 'tmp'

    def open(self, filehandle):
        if 'pdf' in filehandle:
            self.fh = filehandle
            self.outputname = filehandle.split('.')[0]
        else:
            raise ValueError('filehandle invalid')
        with open(filehandle, 'rb') as pdfobj:
            self.pdfReader = PyPDF2.PdfFileReader(filehandle)
        return self.pdfReader

    def decodePages(self):
        if os.path.isfile(os.path.join(self.tempPath, self.outputname + '_decoded.txt')):
            with open(os.path.join(self.tempPath, self.outputname + '_decoded.txt')) as f:
                txt = f.read()
            self.pages = txt
        else:
            pages = []
            for i in range(self.pdfReader.numPages):
                content = self.pdfReader.getPage(i).extractText().encode(
                    'gbk', 'ignore').decode('gbk', 'ignore')
                pages.append(content)
            print(len(self.pdfReader.getPage(i).extractText().encode(
                'gbk', 'ignore').decode('gbk', 'ignore')))
            self.pages = pages
            print(len(pages))
            self.writeTxt(outPath=self.tempPath,
                          filename=self.outputname + '_decoded.txt', translated=False)
        return pages

    def washData(self):
        if os.path.isfile(os.path.join(self.tempPath, self.outputname + '_washed.txt')):
            with open(os.path.join(self.tempPath, self.outputname + '_washed.txt')) as f:
                txt = f.read()
            self.pages = txt
        else:
            _ = []
            for page in self.pages:
                # . \n marked by ####
                newpage = re.sub(r'[\.] *?\n', '####', page)
                # \S\n\S replaced by \S\S5
                newpage = re.sub(r'(\S)\n(\S)', r'\1' + r'\2', newpage)
                # , \n|; \n replaced by , |;
                newpage = re.sub(r'([,|;] )\n', r'\1', newpage)
                # \s
                # newpage = re.sub(r'\s+?\n', ' ', newpage)
                # newpage = re.sub(r'\n\s+?', ' ', newpage)
                newpage = re.sub(r'\n', ' ', newpage)
                newpage = newpage.replace('####', '.\n')
                _.append(newpage)
            self.pages = _
            self.writeTxt(outPath=self.tempPath,
                          filename=self.outputname + '_washed.txt', translated=False)
        return _

    def autoTranslate(self):
        _ = []
        trans = translator.Translator()
        for page in self.pages:
            if not re.findall(r'[\u4e00-\u9fff]+', page):
                _.append(trans.translate(page))
        self.translatedPages = _
        return self.translatedPages

    def translate(self):
        _ = []
        trans = translator.Translator()
        for page in self.pages:
            _.append(trans.translate(page))
        self.translatedPages = _
        return self.translatedPages

    def translatedExist(self):
        _translatedExist = True
        try:
            self.translatedPages
        except AttributeError:
            _translatedExist = False
        return _translatedExist

    def writeTxt(self, outPath='.', filename='output.txt', translated=True):
        if self.outputname:
            filename = self.outputname + '.txt'
        lines = ''

        if translated and self.translatedExist():
            pages = self.translatedPages
        else:
            pages = self.pages
        for page in pages:
            lines += page
        with open(os.path.join(outPath, filename), 'w') as f:
            f.write(lines)

    def abstract(self, outPath='.', filename='abstract.docx', translated=True):
        if self.outputname:
            filename = self.outputname + '_abstract.docx'
        doc_new = docx.Document()
        txt = ''
        for i in self.pages:
            txt += i
        summer = summarizer.StatisticalSummarizer()
        summary = summer.summarize(txt)

        for k in summary.keys():
            summaryTxt = ''
            summaryTxt += str(k) + '\n'
            for i in summary[k]:
                summaryTxt += str(i) + '\n'
            doc_new.add_paragraph(summaryTxt)
        doc_new.save(os.path.join(outPath, filename))

    def writeDocx(self, outPath='.', filename='output.docx', translated=True):
        if self.outputname:
            filename = self.outputname + '.docx'
        doc_new = docx.Document()
        if translated and self.translatedExist():
            pages = self.translatedPages
        else:
            pages = self.pages
        for page in pages:
            for para in page.split('\n'):
                doc_new.add_paragraph(para, style=None)
        doc_new.save(os.path.join(outPath, filename))


if __name__ == '__main__':
    reader = pdfParser('test.pdf')
    reader.decodePages()
    reader.writeTxt(filename='1.txt')
    reader.washData()
    reader.writeTxt(filename='2.txt')
    reader.translate()

    reader.writeTxt(filename='3.txt')
    reader.writeDocx()
    reader.abstract()
